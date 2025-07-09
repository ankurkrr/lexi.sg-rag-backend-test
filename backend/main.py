from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from vector_store import PDFVectorStore
from rag_pipeline import RAGPipeline
import os

app = FastAPI()

class QueryRequest(BaseModel):
    query: str

# Set your PDF/DOCX directory here
PDF_DIR = os.environ.get("PDF_DIR", r"C:\Users\iaman\Vscode Pycharm\RAG Lexi - Test\pdfs")  # Default to 'pdfs' folder in project root

# Initialize FAISS vector store for PDFs
doc_vector_store = PDFVectorStore(PDF_DIR,)

# Initialize RAG pipeline
rag_pipeline = RAGPipeline(doc_vector_store)

@app.post("/query")
async def query_endpoint(request: QueryRequest):
    result = rag_pipeline.run(request.query)
    if not result["answer"]:
        raise HTTPException(status_code=404, detail="No relevant legal text found.")
    return result

# ---
# Sample legal document loader (for demo/testing)
# Place your sample PDFs/DOCX in the 'pdfs/' folder or set PDF_DIR env variable.
# Example script to create a sample DOCX:
if __name__ == "__main__":
    from docx import Document
    os.makedirs(PDF_DIR, exist_ok=True)
    doc = Document()
    doc.add_heading('Sample Legal Document', 0)
    doc.add_paragraph('This is a sample legal document for testing the RAG pipeline.')
    doc.add_paragraph('without a permit is an infraction and insurer is not liable. 18. In Lakhmi Chand v. Reliance General Insurance, (2016) 3 SCC 100, the Court was concerned with an order passed by the National Consumer Disputes Redressal Commission (NCDRC) that had declined the relief to the petitioner therein. The insurer in the said case had taken the plea that the complainant had violated the terms and conditions of the policy, for five passengers were travelling in the goods carrying vehicle at the time of the accident, whereas the permitted seating capacity of the motor vehicle of the appellant was only 1 + 1. The two-Judge Bench referred to Oriental Insurance Co. Ltd. v. Meena Variyal and others, (2007) 5 SCC 428 and expressed the view that in order to avoid liability, the insurer must establish that there was breach on the part of the insured.')
    doc.save(os.path.join(PDF_DIR, 'sample_legal.docx'))
